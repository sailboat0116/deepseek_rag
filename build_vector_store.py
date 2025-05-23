import os
import faiss
import pickle
from docx import Document
from sentence_transformers import SentenceTransformer

def read_docx(path):
    doc = Document(path)
    texts = []

    # 正文段落
    for p in doc.paragraphs:
        if p.text.strip():
            texts.append(p.text.strip())

    # 表格內容（可選）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text.strip())

    print(f"📄 讀取 {os.path.basename(path)}，總段落數：{len(texts)}")
    return "\n".join(texts)

def chunk_text(text, max_length=500):
    sentences = text.replace("\n", " ").split(". ")
    chunks, current = [], ""
    for s in sentences:
        if len(current) + len(s) < max_length:
            current += s + ". "
        else:
            chunks.append(current.strip())
            current = s + ". "
    if current:
        chunks.append(current.strip())
    return chunks

def build_faiss_index(chunks, model_name="all-MiniLM-L6-v2"):
    model = SentenceTransformer(model_name)
    embeddings = model.encode(chunks)
    index = faiss.IndexFlatL2(embeddings.shape[1])
    index.add(embeddings)
    return index, chunks

if __name__ == "__main__":
    folder = r"C:\Users\sailboat\data"
    all_chunks = []

    for filename in os.listdir(folder):
        if filename.lower().endswith(".docx"):
            path = os.path.join(folder, filename)
            try:
                text = read_docx(path)
                chunks = chunk_text(text)
                all_chunks.extend(chunks)
            except Exception as e:
                print(f"❌ 無法讀取 {filename}：{e}")

    print(f"\n🧩 全部文件合併後，總共分段：{len(all_chunks)}")

    if not all_chunks:
        print("❌ 沒有有效段落，請檢查 Word 文件是否有內容。")
        exit()

    index, docs = build_faiss_index(all_chunks)
    faiss.write_index(index, "vector.index")
    with open("docs.pkl", "wb") as f:
        pickle.dump(docs, f)

    print("✅ 向量資料庫已建立完成！")
